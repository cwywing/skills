# -*- coding: utf-8 -*-
"""
将 md素材/ 目录下的鉴别材料 Markdown 文件合并转换为 2 个 Word (.docx) 文档:
  - 文档鉴别材料.docx   (由 document_identification-*.md 合并)
  - 程序鉴别材料.docx   (由 program_identification-*.md 合并)

用法:
    python md2docx.py --workdir docs/软著登记申请

特性:
    - 自动从工作目录 基础资料.md 读取软件名称和版本号，写入页眉
    - 支持标题、段落、代码块、表格、列表、粗体/斜体
    - A4 页面，合理字体和行距
"""

import os
import re
import glob
import sys

_SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, _SCRIPTS_DIR)

from _paths import parse_workdir_arg, load_basic_info  # noqa: E402

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 由 main() 初始化
MATERIAL_DIR = ""
BASIC_INFO_FILE = ""

# 字体配置
FONT_CN = "宋体"
FONT_CODE = "Consolas"
FONT_EN = "Times New Roman"

# 页面边距 (厘米) — 收窄边距增加每页容量
MARGIN_TOP = 2.0
MARGIN_BOTTOM = 2.0
MARGIN_LEFT = 2.5
MARGIN_RIGHT = 2.5

# 字号配置（pt）— 程序鉴别材料用小字
FONT_SIZE_NORMAL = 9
FONT_SIZE_CODE = 8
FONT_SIZE_TABLE = 8
FONT_SIZE_HEADING_BASE = 14

# ===== 合并组定义 =====
# 每组: (输出文件名, md 文件匹配模式列表, 按数字排序)
MERGE_GROUPS = [
    {
        "name": "文档鉴别材料",
        "output": "文档鉴别材料.docx",
        "pattern": "document_identification-*.md",
    },
    {
        "name": "程序鉴别材料",
        "output": "程序鉴别材料.docx",
        "pattern": "program_identification-*.md",
    },
]


def load_software_info():
    """从基础资料.md 读取软件全称和版本号"""
    _, info = load_basic_info(BASIC_INFO_FILE)
    return {
        "全称": info.get("软件全称", ""),
        "版本号": info.get("版本号", ""),
    }


def add_page_header(doc, text):
    """添加页眉"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_page_number(doc):
    """添加页脚页码"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._element.append(fldChar2)


def setup_page(doc):
    """设置 A4 页面和边距"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(MARGIN_TOP)
    section.bottom_margin = Cm(MARGIN_BOTTOM)
    section.left_margin = Cm(MARGIN_LEFT)
    section.right_margin = Cm(MARGIN_RIGHT)


def set_run_font(run, font_size=12, bold=False, italic=False, is_code=False):
    """统一设置 run 字体"""
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if is_code:
        run.font.name = FONT_CODE
    else:
        run.font.name = FONT_EN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_heading(doc, level, text):
    """添加标题"""
    h = doc.add_heading(text, level=min(level, 4))
    for run in h.runs:
        run.font.name = FONT_EN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_rich_paragraph(doc, line, font_size=FONT_SIZE_NORMAL, is_code=False):
    """支持 **粗体** *斜体* `行内代码` 的段落"""
    p = doc.add_paragraph()
    pattern = r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)"
    parts = re.split(pattern, line)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, font_size, bold=True, is_code=is_code)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            set_run_font(run, font_size, italic=True, is_code=is_code)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, font_size, is_code=True)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            run = p.add_run(part)
            set_run_font(run, font_size, is_code=is_code)
    fmt = p.paragraph_format
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(1)
    fmt.line_spacing = Pt(font_size * 1.5)
    return p


def add_table(doc, header_line, rows):
    """添加表格"""
    cols = [c.strip() for c in header_line.strip("|").split("|")]
    col_count = len(cols)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(col_name)
        set_run_font(run, font_size=10, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F2F2")
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)

    for r_idx, row_line in enumerate(rows):
        cells = [c.strip() for c in row_line.strip("|").split("|")]
        for c_idx in range(col_count):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = cells[c_idx] if c_idx < len(cells) else ""
            run = p.add_run(text)
            set_run_font(run, font_size=10)

    doc.add_paragraph()


def add_code_block(doc, lines):
    """添加代码块（灰色背景段落）"""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, font_size=FONT_SIZE_CODE, is_code=True)
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = Pt(10)
        fmt.left_indent = Cm(0.3)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F8F8F8")
        p._element.get_or_add_pPr().append(shd)


def _flush_table(doc, table_lines):
    """将收集的表格行写入文档"""
    if len(table_lines) < 1:
        return
    header = table_lines[0]
    rows = table_lines[1:]
    add_table(doc, header, rows)
    table_lines.clear()


def _add_rich_runs(paragraph, text, font_size=12, is_code=False):
    """向段落添加富文本 runs"""
    pattern = r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, font_size, bold=True, is_code=is_code)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font_size, italic=True, is_code=is_code)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font_size, is_code=True)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, font_size, is_code=is_code)


def _parse_lines_into_doc(doc, lines):
    """将 markdown 行列表解析并追加到已有 Document 对象中"""
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # --- 代码块 ---
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                if table_lines:
                    _flush_table(doc, table_lines)
                    table_lines = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- 表格 ---
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if table_lines:
                _flush_table(doc, table_lines)
                table_lines = []

        # --- 空行 ---
        if not stripped:
            i += 1
            continue

        # --- 标题 ---
        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if m:
                level = len(m.group(1))
                heading_text = m.group(2).strip()
                add_heading(doc, level, heading_text)
                i += 1
                continue

        # --- 水平线 ---
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            set_run_font(run, font_size=8)
            run.font.color.rgb = RGBColor(192, 192, 192)
            i += 1
            continue

        # --- 列表项 ---
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            content = list_match.group(3)
            p = doc.add_paragraph()
            run = p.add_run("  " * indent_level + ("• " if list_match.group(2) in "-*+" else list_match.group(2) + " "))
            set_run_font(run, font_size=FONT_SIZE_NORMAL)
            _add_rich_runs(p, content, font_size=FONT_SIZE_NORMAL)
            fmt = p.paragraph_format
            fmt.space_before = Pt(1)
            fmt.space_after = Pt(1)
            fmt.line_spacing = Pt(13)
            i += 1
            continue

        # --- 普通段落 ---
        add_rich_paragraph(doc, stripped, font_size=FONT_SIZE_NORMAL)
        i += 1

    # 收尾
    if table_lines:
        _flush_table(doc, table_lines)


def _extract_sort_number(filename):
    """从文件名提取数字用于排序，如 program_identification-3.md -> 3"""
    m = re.search(r"-(\d+)\.md$", filename)
    return int(m.group(1)) if m else 0


def _read_md_lines(md_path):
    """读取 md 文件，返回行列表"""
    for enc in ("utf-8", "utf-8-sig", "gbk"):
        try:
            with open(md_path, "r", encoding=enc) as f:
                return f.readlines()
        except UnicodeDecodeError:
            continue
    return []


def merge_and_convert(group, software_info):
    """将一组 md 文件合并转换为一个 docx"""
    # 按数字排序收集 md 文件
    md_files = sorted(glob.glob(os.path.join(MATERIAL_DIR, group["pattern"])),
                      key=lambda p: _extract_sort_number(os.path.basename(p)))

    if not md_files:
        print("  [跳过] 未找到匹配文件: {}".format(group["pattern"]))
        return

    print("  包含文件:")
    for f in md_files:
        print("    + {}".format(os.path.basename(f)))

    # 创建文档
    doc = Document()
    setup_page(doc)

    header_text = "{} {}".format(software_info["全称"], software_info["版本号"])
    add_page_header(doc, header_text)
    add_page_number(doc)

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(FONT_SIZE_NORMAL)
    style.paragraph_format.line_spacing = Pt(13)

    # 逐个文件追加内容
    for md_path in md_files:
        lines = _read_md_lines(md_path)
        if lines:
            _parse_lines_into_doc(doc, lines)

    # 保存
    output_path = os.path.join(MATERIAL_DIR, group["output"])
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print("  -> 已生成: {} ({}KB)".format(group["output"], size_kb))


def main():
    global MATERIAL_DIR, BASIC_INFO_FILE

    paths_meta = parse_workdir_arg()
    MATERIAL_DIR = paths_meta["material_dir"]
    BASIC_INFO_FILE = paths_meta["basic_info"]

    print("=" * 50)
    print("  MD -> Word 合并转换工具")
    print("  同组 MD 合并为一个 DOCX")
    print("  工作目录: {}".format(paths_meta["workdir"]))
    print("=" * 50)

    info = load_software_info()
    print("\n软件信息: {} {}".format(info["全称"], info["版本号"]))
    print()

    for group in MERGE_GROUPS:
        print("[转换] {} -> {}".format(group["name"], group["output"]))
        try:
            merge_and_convert(group, info)
        except Exception as e:
            print("  [错误] {}: {}".format(group["name"], e))
        print()

    print("=" * 50)
    print("  全部转换完成!")
    print("=" * 50)


if __name__ == "__main__":
    main()
